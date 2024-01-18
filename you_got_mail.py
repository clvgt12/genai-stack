#!/usr/bin/env python

#%pip install langchain langchain-community langchainhub langchain-openai neo4j tiktoken unstructured "unstructured[all-docs]" pdfplumber  beautifulsoup4 lxml python-docx


from langchain import hub
from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain_community.document_loaders import UnstructuredEmailLoader
from langchain_community.vectorstores import Neo4jVector
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.embeddings import SentenceTransformerEmbeddings


# Define Documents and metadata objects

class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        return f"[Document(page_content={self.page_content}, metadata={self.metadata})]"


# Define Password Store class to interact with Unix utility password-store "pass"

import os
import subprocess

class PasswordStore():

    def show(self, token: str, **kwargs) -> None:
        try:
            # Run the pass command and capture the output
            result = subprocess.run(['pass', 'show', token], 
                                    check=True, 
                                    stdout=subprocess.PIPE, 
                                    stderr=subprocess.PIPE, 
                                    text=True)
            
            # Extract the API key from the output
            result = result.stdout.strip()
            return result
        except subprocess.CalledProcessError as e:
            # Handle errors if the command fails
            print(f"PasswordStore.show(): Error occurred: {e.stderr}")
            return None

# Use the function to get the API key
pwdstore = PasswordStore()
os.environ["OPENAI_API_KEY"] = pwdstore.show('OpenAI/openai_api_key')


# ```
# docker run \
#     --name neo4j \
#     -p 7474:7474 -p 7687:7687 \
#     -d \
#     -e NEO4J_AUTH=neo4j/password \
#     neo4j:5.11
# ```

# Neo4jVector requires the Neo4j database credentials

# url = "bolt://localhost:7687"
# username = "neo4j"
# password = "password"

# You can also use environment variables instead of directly passing named parameters
os.environ["NEO4J_URI"] = "bolt://localhost:7687"
os.environ["NEO4J_USERNAME"] = "neo4j"
os.environ["NEO4J_PASSWORD"] = "password"


# Store processed email file in the vector store

# The Neo4jVector Module will connect to Neo4j and create a vector index if needed.

def store_email_file_in_vectorstore(texts,url=os.environ["NEO4J_URI"],username=os.environ["NEO4J_USERNAME"],password=os.environ["NEO4J_PASSWORD"],destroy_information_flag=False):
        vectorstore = Neo4jVector.from_texts(
                texts=texts, 
                embedding=SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2"), 
                url=url, 
                username=username, 
                password=password, 
                index_name="SRRP_email", 
                node_label="SRRP_email",
                pre_delete_collection=destroy_information_flag  # Delete existing data? True or False
        )
        return(vectorstore)


# Return a reference to the existing vectorstore

def get_vectorstore(url=os.environ["NEO4J_URI"],username=os.environ["NEO4J_USERNAME"],password=os.environ["NEO4J_PASSWORD"],destroy_information_flag=False):
        vectorstore = Neo4jVector.from_existing_index(
                embedding=SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2"), 
                url=url, 
                username=username, 
                password=password, 
                index_name="SRRP_email", 
                node_label="SRRP_email",
                pre_delete_collection=destroy_information_flag  # Delete existing data? True or False
        )
        return(vectorstore)


# Load, parse and split specified email file
# 
# First try to process any attachments.  If an error is trapped, then process without

import email
from email.parser import BytesParser
from bs4 import BeautifulSoup
import docx
import json
import pdfplumber
import io
import re

def clean_text(text):
    return(re.sub(r'[^(\w\s)]', '', text))

def docx2text(docx_bytes):
    """
    Extracts and returns the text multipart_content from a DOCX file.

    :param docx_bytes: Bytes content of a DOCX file.
    :return: Plain text extracted from the DOCX file.
    """
    doc = docx.Document(io.BytesIO(docx_bytes))
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    return clean_text('\n'.join(full_text))

def pdf2text(pdf_bytes):
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
        return clean_text(text)

def html2text(html_content):
    """
    Extracts and returns the text content from HTML.

    :param html_content: String containing HTML content.
    :return: Plain text extracted from HTML.
    """
    soup = BeautifulSoup(html_content, 'lxml')
    return clean_text(soup.get_text())

def filter_email_parts(eml_file_path):
    """
    Reads an EML file, filters out parts not matching specified MIME types.

    :param eml_file_path: Path to the .eml file
    :return: A list of text strings that represent the text contents of the email file
    """
    with open(eml_file_path, 'rb') as f:
        msg = BytesParser().parse(f)

    content = []
    multipart_content = []

    if msg.is_multipart():
        for part in msg.walk():
            mime_type=part.get_content_type()
            if mime_type == "application/pdf":
                pdf_bytes = part.get_payload(decode=True)
                pdf_text = pdf2text(pdf_bytes)
                multipart_content.append(pdf_text)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                docx_bytes = part.get_payload(decode=True)
                word_text = docx2text(docx_bytes)
                multipart_content.append(word_text)
            elif mime_type == "text/html":
                html_content = part.get_payload(decode=True).decode('utf-8')
                text_content = html2text(html_content)
                content.append(text_content)
            elif mime_type == "text/plain":
                multipart_content.append(clean_text(part.get_payload(decode=True).decode('utf-8')))
    else:
        mime_type=msg.get_content_type()
        if mime_type == "text/html":
            html_content = msg.get_payload(decode=True).decode('utf-8')
            text_content = html2text(html_content)
            content.append(text_content)
        elif mime_type == "text/plain":
            content.append(clean_text(msg.get_payload(decode=True).decode('utf-8')))

    content = multipart_content if multipart_content != [] else content
    return(' '.join(content))

def process_email_file(email_filename):
    texts = filter_email_parts(email_filename)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    docs = text_splitter.split_text(texts)
    return(docs)

# Connect to IMAP email server, download email from SRRP team members

import imaplib
import datetime

def dump_email_to_file(imap_server, email_account, password, from_addresses, days):

    tmp_filepath="/tmp/SRRP_email"
    destroy_information_flag = True

    # Connect to IMAP server
    mail = imaplib.IMAP4_SSL(imap_server)
    mail.login(email_account, password)
    mail.select('inbox')

    # Calculate the date X days ago
    date_days_ago = (datetime.datetime.now() - datetime.timedelta(days=days)).strftime("%d-%b-%Y")

    mail_ids = []
    for address in from_addresses:
        # Search emails from a specific address since the date
        _, data = mail.search(None, f'(FROM "{address}" SINCE "{date_days_ago}")')
        mail_ids.extend(data[0].split())

    total_emails = len(mail_ids)
    print(f"Total emails to process: {total_emails}")

    # Process each email
    for i, num in enumerate(mail_ids, 1):  # Start counting from 1
        _, data = mail.fetch(num, '(RFC822)')
        raw_email = data[0][1]

        # Write each email to a separate file
        email_filename = f"{tmp_filepath}_{num.decode('utf-8')}.eml"
        with open(email_filename, 'wb') as email_file:
            email_file.write(raw_email)

        docs = process_email_file(email_filename)
        os.remove(email_filename)
        vectorstore = store_email_file_in_vectorstore(docs,destroy_information_flag=destroy_information_flag)
        destroy_information_flag = False if destroy_information_flag == True else destroy_information_flag        
        # Print the percent complete
        percent_complete = (i / total_emails) * 100
        if i % 10 == 0 or i == total_emails:
            print(f"Processed {i}/{total_emails} emails ({percent_complete:.2f}% complete)")

    # Close connections
    mail.logout()

    # Return reference to the vectorstore
    return(vectorstore)

# Usage example
imap_password = pwdstore.show('Email/cvitalos@yahoo.com/imap')
from_addresses = ["guitarduo@verizon.net", "WardenA68@yahoo.com", "jamesmespo@gmail.com", "sahare0515@gmail.com", "julia@njhighlandscoalition.org", 
                    "elliot@njhighlandscoalition.org", "kdolsky@optonline.net", "ssolaun@gmail.com", "bobmossnj@verizon.net"]
vectorstore = dump_email_to_file('imap.mail.yahoo.com', 'cvitalos@yahoo.com', imap_password, from_addresses, 90)

# Now use a prompting technique, and feed the result through a LLM

from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda, RunnablePassthrough

retriever = vectorstore.as_retriever()

template = """Answer the question based only on the following context:
{context}

Question: {question}
"""
prompt = ChatPromptTemplate.from_template(template)

model = ChatOpenAI()

chain = (
    {"context": retriever, "question": RunnablePassthrough()}
    | prompt
    | model
    | StrOutputParser()
)

# Conversational Retrieval Chain
# 
# Add in conversation history. This primarily means adding in chat_message_history
# 

from operator import itemgetter

from langchain.schema import format_document
from langchain_core.messages import AIMessage, HumanMessage, get_buffer_string
from langchain_core.runnables import RunnableParallel

from langchain.prompts.prompt import PromptTemplate

_template = """Given the following conversation and a follow up question, rephrase the follow up question to be a standalone question, in its original language.

Chat History:
{chat_history}
Follow Up Input: {question}
Standalone question:"""
CONDENSE_QUESTION_PROMPT = PromptTemplate.from_template(_template)

template = """Answer the question based only on the following context:
{context}

Question: {question}
"""
ANSWER_PROMPT = ChatPromptTemplate.from_template(template)

DEFAULT_DOCUMENT_PROMPT = PromptTemplate.from_template(template="{page_content}")

def _combine_documents(
    docs, document_prompt=DEFAULT_DOCUMENT_PROMPT, document_separator="\n\n"
):
    doc_strings = [format_document(doc, document_prompt) for doc in docs]
    return document_separator.join(doc_strings)

_inputs = RunnableParallel(
    standalone_question=RunnablePassthrough.assign(
        chat_history=lambda x: get_buffer_string(x["chat_history"])
    )
    | CONDENSE_QUESTION_PROMPT
    | ChatOpenAI(temperature=0)
    | StrOutputParser(),
)
_context = {
    "context": itemgetter("standalone_question") | retriever | _combine_documents,
    "question": lambda x: x["standalone_question"],
}
# conversational_qa_chain = _inputs | _context | ANSWER_PROMPT | ChatOpenAI()

# conversational_qa_chain.invoke(
#     {
#         "question": query,
#         "chat_history": [],
#     }
# )

# With Memory and returning source documents

from operator import itemgetter

from langchain.memory import ConversationBufferMemory

memory = ConversationBufferMemory(
    return_messages=True, output_key="answer", input_key="question"
)

# First we add a step to load memory
# This adds a "memory" key to the input object
loaded_memory = RunnablePassthrough.assign(
    chat_history=RunnableLambda(memory.load_memory_variables) | itemgetter("history"),
)
# Now we calculate the standalone question
standalone_question = {
    "standalone_question": {
        "question": lambda x: x["question"],
        "chat_history": lambda x: get_buffer_string(x["chat_history"]),
    }
    | CONDENSE_QUESTION_PROMPT
    | ChatOpenAI(temperature=0)
    | StrOutputParser(),
}
# Now we retrieve the documents
retrieved_documents = {
    "docs": itemgetter("standalone_question") | retriever,
    "question": lambda x: x["standalone_question"],
}
# Now we construct the inputs for the final prompt
final_inputs = {
    "context": lambda x: _combine_documents(x["docs"]),
    "question": itemgetter("question"),
}
# And finally, we do the part that returns the answers
answer = {
    "answer": final_inputs | ANSWER_PROMPT | ChatOpenAI(),
    "docs": itemgetter("docs"),
}
# And now we put it all together!
final_chain = loaded_memory | standalone_question | retrieved_documents | answer

def ask_question(question, final_chain):
    inputs = {"question": question}
    result = final_chain.invoke(inputs)
    return(result)

def save_memory(memory, question, result):
    memory.save_context({"question": question}, {"answer": result["answer"].content})
    memory.load_memory_variables({})
    return(memory)

question = "Describe what is known about the October 27th meeting"
result = ask_question(question, final_chain)
memory = save_memory(memory, question, result)
print(result)

question = "Identify who attended the Oct 27th meeting, and who did not"
result = ask_question(question, final_chain)
memory = save_memory(memory, question, result)
print(result)

question = "Describe what was the outcome taken by the commissioners during the October 27th meeting"
result = ask_question(question, final_chain)
memory = save_memory(memory, question, result)
print(result)

question = "Identify any discussion in the email threads regarding logging at Merrill Creek - also known as MC"
result = ask_question(question, final_chain)
memory = save_memory(memory, question, result)
print(result)