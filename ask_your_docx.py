import os
import io

import streamlit as st
from langchain.chains import RetrievalQA
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.callbacks.base import BaseCallbackHandler
from langchain_community.vectorstores import Neo4jVector
from streamlit.logger import get_logger
from chains import (
    load_embedding_model,
    load_llm,
    lc_configure_qa_rag_chain,
)

from docx import Document
from odf import teletype
from odf.table import Table, TableRow, TableCell
from odf.text import P
from odf.opendocument import load as load_odf

# load api key lib
from dotenv import load_dotenv

load_dotenv(".env")


url = os.getenv("NEO4J_URI")
username = os.getenv("NEO4J_USERNAME")
password = os.getenv("NEO4J_PASSWORD")
ollama_base_url = os.getenv("OLLAMA_BASE_URL")
embedding_model_name = os.getenv("EMBEDDING_MODEL")
llm_name = os.getenv("LLM")
# Remapping for Langchain Neo4j integration
os.environ["NEO4J_URL"] = url

logger = get_logger(__name__)


embeddings, dimension = load_embedding_model(
    embedding_model_name, config={"ollama_base_url": ollama_base_url}, logger=logger
)


class StreamHandler(BaseCallbackHandler):
    def __init__(self, container, initial_text=""):
        self.container = container
        self.text = initial_text

    def on_llm_new_token(self, token: str, **kwargs) -> None:
        self.text += token
        self.container.markdown(self.text)


llm = load_llm(llm_name, logger=logger, config={"ollama_base_url": ollama_base_url})


def main():
    st.header("📄Chat with your document")

    # upload a your pdf file
    doc = st.file_uploader("Upload your PDF, DOCX, ODT or TXT:", type=["pdf", "docx", "odt", "txt", "log"])

    if doc is not None:

        file_type = doc.name.split(".")[-1].lower()
        if file_type == 'pdf':
            # Read in Adobe PDF files
            pdf_reader = PdfReader(doc)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file_type == 'docx':
            # Read in MS Word files
            msword = Document(io.BytesIO(doc.read()))
            text = '\n'.join([paragraph.text for paragraph in msword.paragraphs])
            # Iterate over each table in the document
            for table in msword.tables:
                headers = []  # To store the feature labels (FL1, FL2, ..., FLN)
                category_label = ""  # To store the category label (CL)
                sentences = []  # To store the constructed sentences
                for row_index, row in enumerate(table.rows):
                    cell_texts = [cell.text.strip() for cell in row.cells]
                    if row_index == 0:
                        category_label = cell_texts[0]  # The first cell of the first row is the category label
                        headers = cell_texts[1:]  # The rest of the first row are the feature labels
                    else:
                        example_label = cell_texts[0]  # The first cell of each subsequent row is the example label
                        features = cell_texts[1:]  # The rest are the feature values
                        # Construct the sentence for the current row
                        sentence_parts = [f"{category_label} {example_label} has the following characteristic(s) "]
                        sentence_parts.extend([f"{headers[idx]}: {features[idx]}" for idx in range(len(features))])
                        sentence = ", and ".join(sentence_parts) + "."
                        sentence = sentence.replace("has the following characteristic(s) , and", "has the following characteristic(s):")
                        sentences.append(sentence)
                text += " ".join(sentences) + "\n\n"  # Add two newlines after each table for readability
        elif file_type == 'odt':
            # Read in Libre Office Writer ODT files
            odt_doc = load_odf(io.BytesIO(doc.read()))
            paragraphs = odt_doc.getElementsByType(P)
            text = '\n'.join(teletype.extractText(p) for p in paragraphs)
            # Find all tables in the document
            tables = odt_doc.getElementsByType(Table)
            # Iterate over each table
            for table in tables:
                headers = []  # To store the feature labels (FL1, FL2, ..., FLN)
                category_label = ""  # To store the category label (CL)
                sentences = []  # To store the constructed sentences
                # Iterate over the table's rows
                for row_index, row in enumerate(table.getElementsByType(TableRow)):
                    row_cells = row.getElementsByType(TableCell)
                    cell_texts = [teletype.extractText(cell.getElementsByType(P)[0]) if cell.getElementsByType(P) else "" for cell in row_cells]
                    if row_index == 0:
                        category_label = cell_texts[0]  # The first cell of the first row is the category label
                        headers = cell_texts[1:]  # The rest of the first row are the feature labels
                    else:
                        example_label = cell_texts[0]  # The first cell of each subsequent row is the example label
                        features = cell_texts[1:]  # The rest are the feature values
                        # Construct the sentence for the current row
                        sentence_parts = [f"{category_label} {example_label} has the following characteristic(s) "]
                        sentence_parts.extend([f"{headers[idx]}: {feat}" for idx, feat in enumerate(features)])
                        sentence = ", and ".join(sentence_parts) + "."
                        sentence = sentence.replace("has the following characteristic(s) , and", "has the following characteristic(s):")
                        sentences.append(sentence)
                text += " ".join(sentences) + "\n\n"  # Add two newlines after each table for readability
        elif file_type in ['txt', 'log']:
            # Read in plain text files
            text = doc.read().decode('utf-8')

        # langchain_textspliter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200, length_function=len
        )

        chunks = text_splitter.split_text(text=text)

        # Store the chunks part in db (vector)
        vectorstore = Neo4jVector.from_texts(
            chunks,
            url=url,
            username=username,
            password=password,
            embedding=embeddings,
            index_name="ask_your_docx",
            node_label="ask_your_docx",
            pre_delete_collection=True,  # Delete existing data
        )

        qa = lc_configure_qa_rag_chain(llm, embeddings, url, username, password, "ask_your_docx")

        # Accept user questions/query
        query = st.text_input("Ask questions about your document")

        if query:
            stream_handler = StreamHandler(st.empty())
            logger.info(f"Query: {query}")
            response = qa.run(query, callbacks=[stream_handler])
            logger.info(f"Response: {response}")


if __name__ == "__main__":
    main()
